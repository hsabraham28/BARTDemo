from zipfile import ZipFile
from docx import Document
from IPython.display import FileLink
import pandas as pd

def read_word_doc_zip_as_df(zipFile):
    dfs = []
    
    with ZipFile(zipFile, 'r') as zipObj:
        # Get list of files names in zip
        listOfiles = zipObj.namelist()
        # Iterate over the list of file names in given list & print them
        for elem in listOfiles:
            file = zipObj.open(elem)
            wordDoc = Document(file)
            for table in wordDoc.tables:
                data = [[cell.text for cell in row.cells] for row in table.rows]
                dfs.append(pd.DataFrame(data))
    
    df = pd.concat(dfs, ignore_index=True)
    #column labels are the values of the first row
    df.columns = df.iloc[0] 
    df = df[1:]
    
    print('done')
    
    return df
    
    
def word_doc_to_df(fileName):
    wordDoc = Document(fileName)  
    dfs = []
    for table in wordDoc.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        dfs.append(pd.DataFrame(data))
    #column labels are the values of the first row
    df = pd.concat(dfs, ignore_index=True)
    df.columns = df.iloc[0] 
    df = df[1:]
    print('done')
    return df
        
    
def download_df_as_csv(df, filename):
    from IPython.display import HTML
    import base64
    def create_download_link(df, title = "Download CSV file", filename = filename):  
        csv = df.to_csv()
        b64 = base64.b64encode(csv.encode())
        payload = b64.decode()
        html = '<a download="{filename}" href="data:text/csv;base64,{payload}" target="_blank">{title}</a>'
        html = html.format(payload=payload,title=title,filename=filename)
        return HTML(html)
    
    create_download_link(occ_df)